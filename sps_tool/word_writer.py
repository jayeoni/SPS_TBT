"""
Bilingual Word document generator.
Translates all 13 rows of the WTO SPS notification form following the
exact format used in reference translated files (26.3월 SPS 통보문_금영★).
"""
import re
import shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

KOREAN_FONT = '맑은 고딕'
LIME_RGB    = (204, 255, 153)

LANG_KR = {
    'english': '영어', 'spanish': '스페인어', 'portuguese': '포르투갈어',
    'french': '프랑스어', 'arabic': '아랍어', 'chinese': '중국어',
    'russian': '러시아어', 'german': '독일어', 'japanese': '일본어',
}

MONTH_KR = {
    'january': '1월', 'february': '2월', 'march': '3월', 'april': '4월',
    'may': '5월', 'june': '6월', 'july': '7월', 'august': '8월',
    'september': '9월', 'october': '10월', 'november': '11월', 'december': '12월',
}

# Row detection patterns: checked against first 150 chars of content cell (lowercase)
ROW_PATTERNS = {
    'notifying_member': ['notifying member', '통보회원국'],
    'agency':           ['agency responsible', '담당기관'],
    'products':         ['products covered', '적용대상품목'],
    'regions':          ['regions or countries', 'regions/countries', '지역/국가', '국가/지역'],
    'title':            ['title of the notified', 'title of the proposed', '제목'],
    'description':      ['description of content', '내용의 설명'],
    'objective':        ['objective and rationale'],
    'standards':        ['is there a relevant international standard'],
    'urgent_reason':    ['nature of the urgent problem'],
    'other_docs':       ['other relevant documents'],
    'adoption_date':    ['proposed date of adoption'],
    'entry_force':      ['proposed date of entry into force'],
    'comments':         ['final date for comments'],
    'texts_available':  ['text(s) available from', 'texts available from'],
    # Addendum-specific rows
    'addendum_intro':           ['the following communication', 'being circulated at the request'],
    'addendum_country_advises': ['hereby advises', 'hereby notifies'],
    'addendum_concerns':        ['this addendum concerns'],
    'addendum_comment_period_sec': ['comment period:'],
    'addendum_agency_comments': ['agency or authority designated to handle comments'],
}

# Known agency name → Korean translation (checked against original English cell text)
# Used to override unreliable LLM translations for specific agencies.
KNOWN_AGENCIES = [
    (re.compile(r'state phytosanitary service', re.IGNORECASE), '식물위생청(SFE)'),
    (re.compile(r'Institute for Agricultural and Livestock Protection and Health|\bIPSA\b', re.IGNORECASE), '농축산물보호위생청(IPSA)'),
    (re.compile(r'ministry of agriculture and livestock', re.IGNORECASE), '농축산부(MAG)'),
]

OBJECTIVE_OPTIONS = [
    ('food safety',            '식품안전'),
    ('animal health',          '동물위생'),
    ('plant protection',       '식물보호'),
    ('protect humans from',    '동식물 해충/질병으로부터 사람 보호'),
    ('protect territory from', '해충으로 인한 피해로부터의 영토 보호'),
]

# IPPC ISPM titles in Korean (official translations)
ISPM_KR = {
    1:  'SPS 적용 원칙',
    2:  '병해충위험분석 개요',
    3:  '생물학적 방제제의 수출, 운송, 수입 및 방출 지침',
    4:  '병해충발생비지역 설정 요건',
    5:  '식물위생 용어 해설',
    6:  '모니터링 지침',
    7:  '식물위생증명체계',
    8:  '지역의 병해충 발생현황 규명',
    9:  '병해충 박멸 프로그램 지침',
    10: '병해충비발생 생산지 및 생산장소 설정 요건',
    11: '검역병해충에 대한 병해충위험분석',
    12: '식물위생증명서',
    13: '부적합 및 긴급조치의 통보 지침',
    14: '병해충 위험관리를 위한 종합방제체계 지침',
    15: '국제무역에 있어서 목재포장재의 규정',
    16: '규제비검역병해충: 개념 및 적용',
    17: '병해충 보고',
    18: '식물위생조치로서 방사선 조사의 사용 지침',
    19: '규제 병해충 목록 지침',
    20: '식물위생 수입 규정 체계 지침',
    21: '규제비검역병해충의 병해충위험분석',
    22: '병해충 저발생 지역 설정 요건',
    23: '검사 지침',
    24: '동등성 규정 및 인정 지침',
    25: '환적 화물',
    26: '과실파리(Tephritidae)의 병해충발생비지역 설정',
    27: '규제 병해충 진단 규정',
    28: '규제 병해충 식물위생처리',
    29: '병해충발생비지역 및 저발생 지역의 인정',
    30: '과실파리(Tephritidae)의 저발생 지역 설정',
    31: '화물 표본추출 방법론',
    32: '병해충 위험에 따른 품목 분류',
    33: '병해충발생비 감자(Solanum tuberosum) 종자',
    34: '식물 수입 후 격리검역소의 설계 및 운영',
    35: '과실파리 위험관리를 위한 종합방제체계',
    36: '재배용 식물의 종합방제조치',
    38: '종자의 국제이동',
    39: '신선 과실 및 채소의 국제이동',
    40: '재배용 식물과 관련된 재배 배지의 국제이동',
    41: '사용한 차량, 기계 및 장비의 국제이동',
    42: '식물위생조치로서 온도처리 사용 요건',
    43: '목재 식물위생조치로서 열처리 사용 요건',
}


# ── Helpers ──────────────────────────────────────────────────────────────────

def _unique_cells(row):
    seen = set()
    result = []
    for cell in row.cells:
        cid = id(cell._tc)
        if cid not in seen:
            seen.add(cid)
            result.append(cell)
    return result


def _translate_date(text: str) -> str:
    """Convert 'D Month YYYY' patterns to 'YYYY년 M월 D일'."""
    def _repl(m):
        day = str(int(m.group(1)))
        month_kr = MONTH_KR.get(m.group(2).lower(), m.group(2))
        return f'{m.group(3)}년 {month_kr} {day}일'
    return re.sub(
        r'(\d{1,2})\s+'
        r'(January|February|March|April|May|June|July|August|'
        r'September|October|November|December)\s+(\d{4})',
        _repl, text, flags=re.IGNORECASE,
    )


def _translate_date_phrase(text: str) -> str:
    """Translate standard date phrases and apply date conversion."""
    text = re.sub(r'(?i)to be determined after the end of the consultation period\.?',
                  '의견수렴기간 종료 후 결정', text)
    text = re.sub(r'(?i)to be determined', '추후 결정', text)
    text = re.sub(
        r'(?i)approximately\s+(\d+)\s+days?\s+from\s+the\s+date\s+of\s+publication[^\n.]*',
        lambda m: f'관보 게재일로부터 약 {m.group(1)}일 후', text)
    text = re.sub(
        r'(?i)(\d+)\s+days?\s+after\s+publication\s+in\s+the\s+official\s+journal\.?',
        lambda m: f'관보 게재일로부터 {m.group(1)}일 후', text)
    text = re.sub(r'(?i)upon publication in the official journal\.?', '관보게재일', text)
    text = re.sub(r'(?i)the resolution will enter into force upon signature\.?',
                  '본 결의안은 서명 시 발효예정', text)
    text = re.sub(r'(?i)\bupon signature\.?', '서명 시 발효예정', text)
    return _translate_date(text)


def _checkbox(text: str, option_prefix: str) -> str:
    """Return '[X]' or '[  ]' based on whether option is checked in text."""
    m = re.search(r'\[([Xx☒ ]*)\]\s*' + re.escape(option_prefix[:18]),
                  text, re.IGNORECASE)
    if m:
        inner = m.group(1).strip().lower()
        return '[X]' if inner in ('x', '☒') else '[  ]'
    return '[  ]'


def _expand_ispm_numbers(raw_text: str) -> str:
    """Return 'ISPM 제 N장, 제 M장' string parsed from raw_text, or '' if none."""
    numbers = re.findall(r'\b(\d+)\b', raw_text)
    if not numbers:
        return ''
    return 'ISPM ' + ', '.join(f'제 {n}장' for n in numbers)


def _extract_email(text: str) -> str:
    """Extract first email address from text, or ''."""
    m = re.search(r'[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}', text)
    return m.group(0) if m else ''


def _get_cell_style(cell):
    """Return (font_size, para_style, bold, italic, underline) in a single pass.

    bold/italic/underline are set only when uniform across all content runs;
    mixed cells return None for that property so Korean text stays plain.
    """
    font_size = para_style = None
    seen_bold = seen_italic = seen_underline = set()

    for para in cell.paragraphs:
        if para_style is None and para.runs:
            para_style = para.style.name
        for run in para.runs:
            if font_size is None and run.font.size:
                font_size = run.font.size
            if run.text.strip():
                seen_bold.add(run.bold is True)
                seen_italic.add(run.italic is True)
                seen_underline.add(run.underline is True)

    # Apply a property only when every content run agrees (all True or all non-True)
    bold      = True if seen_bold      == {True} else None
    italic    = True if seen_italic    == {True} else None
    underline = True if seen_underline == {True} else None
    return font_size, para_style, bold, italic, underline


def _apply_korean_font(run):
    run.font.name = KOREAN_FONT
    rPr = run._r.get_or_add_rPr()
    for existing in rPr.findall(qn('w:rFonts')):
        rPr.remove(existing)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    KOREAN_FONT)
    rFonts.set(qn('w:hAnsi'),   KOREAN_FONT)
    rFonts.set(qn('w:eastAsia'), KOREAN_FONT)
    rPr.append(rFonts)


def _set_cell_bg(cell, rgb: tuple):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(*rgb)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    tcPr.append(shd)


def _add_paragraph(cell, text: str, font_size=None, style_name=None, bold=None, italic=None, underline=None):
    try:
        para = cell.add_paragraph(style=style_name)
    except Exception:
        para = cell.add_paragraph()
    run = para.add_run(text)
    _apply_korean_font(run)
    if font_size:
        run.font.size = font_size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline


# ── Per-row translation builders ──────────────────────────────────────────────

def _row_notifying_member(cell_text, t):
    country_kr = t.get('통보국_kr', '')
    if not country_kr:
        return []
    return [
        f'통보 회원국: {country_kr}',
        '적용 가능할 경우, 관련 지방 정부의 명칭 기재:',
    ]


def _row_agency(cell_text, t):
    # Check known agencies against original English text; build from hardcoded table if found
    found = [kr for pattern, kr in KNOWN_AGENCIES if pattern.search(cell_text)]
    if found:
        agency_kr = ' / '.join(found)
    else:
        agency_kr = t.get('담당기관_kr', '')
    if not agency_kr:
        return []
    return [f'담당 기관: {agency_kr}']


def _row_products(cell_text, t):
    products_kr = t.get('해당품목', '')
    if not products_kr:
        return []
    return [f'해당 품목(WTO의 국가별 세부 HS코드번호; 가능할 시 ICS 번호 표기): {products_kr}']


def _row_regions(cell_text, t):
    all_cb  = _checkbox(cell_text, 'All trading partners')
    spec_cb = _checkbox(cell_text, 'Specific regions')

    # Spanish fallback: "Todos los interlocutores comerciales" / "Regiones o países específicos"
    if all_cb == '[  ]' and re.search(r'\[x\][^\n]*todos los interlocutores', cell_text, re.IGNORECASE):
        all_cb = '[X]'
    if spec_cb == '[  ]' and re.search(r'\[x\][^\n]*espec[íi]ficos', cell_text, re.IGNORECASE):
        spec_cb = '[X]'

    countries_kr = t.get('해당국가', '')
    spec_text = '특정 지역 및 국가: '
    if spec_cb == '[X]' and countries_kr and countries_kr != '모든 교역국':
        spec_text += countries_kr

    return [
        '영향을 받을 수 있는 지역이나 국가(관련되거나 가능한 한 최대한): ',
        f'{all_cb}\t모든 교역국  ',
        f'{spec_cb}    {spec_text}',
    ]


def _row_title(cell_text, t):
    title_kr = t.get('제목', '')
    if not title_kr:
        return []

    lang_m = re.search(r'Language\(s\):\s*([^\n]+)', cell_text, re.IGNORECASE)
    lang_raw = lang_m.group(1).strip() if lang_m else ''
    lang_raw = re.sub(r'\s+Number of pages[^\n]*', '', lang_raw, flags=re.IGNORECASE).strip()
    if lang_raw:
        parts = re.split(r'\s+and\s+|\s*,\s*|\s*&\s*', lang_raw, flags=re.IGNORECASE)
        kr_parts = [LANG_KR.get(p.strip().lower(), p.strip()) for p in parts if p.strip()]
        lang_kr = ' 및 '.join(kr_parts)
    else:
        lang_kr = ''

    pages_m = re.search(r'Number of pages:\s*(\S+)', cell_text, re.IGNORECASE)
    pages = pages_m.group(1) if pages_m else ''

    line = f'통보 문서의 제목: {title_kr}'
    if lang_kr:
        line += f'  언어: {lang_kr}'

    urls = re.findall(r'https?://\S+', cell_text)
    result = [line] + urls
    if pages:
        result.append(f'페이지수: {pages}')
    return result


def _row_description(cell_text, t):
    desc_kr = t.get('내용', '')
    if not desc_kr:
        return []
    lines = [s.strip() for s in desc_kr.split('\n') if s.strip()]
    if not lines:
        return []
    return [f'내용 설명: {lines[0]}'] + lines[1:]


_OBJECTIVE_MOKJEOK_KEYS = ['식품안전', '동물위생', '식물보호', '사람 보호', '영토 보호']


def _row_objective(cell_text, t):
    mokjeok = re.sub(r'\s+', '', t.get('목적', ''))
    parts = []
    for (eng_prefix, kr_label), key in zip(OBJECTIVE_OPTIONS, _OBJECTIVE_MOKJEOK_KEYS):
        cb = _checkbox(cell_text, eng_prefix)
        if cb == '[  ]' and re.sub(r'\s+', '', key) in mokjeok:
            cb = '[X]'
        parts.append(f'{cb} {kr_label}')
    lines = ['목적 및 근거: ' + ', '.join(parts)]
    if t.get('목적_근거'):
        lines += [s.strip() for s in t['목적_근거'].split('\n') if s.strip()]
    return lines


def _row_standards(cell_text, t):
    codex_cb = _checkbox(cell_text, 'Codex Alimentarius')
    oie_cb   = _checkbox(cell_text, 'World O')  # matches both Organisation/Organization
    ippc_cb  = _checkbox(cell_text, 'International Plant Protection')
    none_cb  = _checkbox(cell_text, 'None')

    # Spanish/Portuguese fallbacks
    if oie_cb == '[  ]' and re.search(r'(?:\[[Xx☒]\]|☒)[^\n]*Organizaci[oó]n Mundial', cell_text, re.IGNORECASE):
        oie_cb = '[X]'
    if ippc_cb == '[  ]' and re.search(r'(?:\[[Xx☒]\]|☒)[^\n]*Convenci[oó]n Internacional|CIPF', cell_text, re.IGNORECASE):
        ippc_cb = '[X]'
    if none_cb == '[  ]' and re.search(r'(?:\[[Xx☒]\]|☒)[^\n]*Ninguna?', cell_text, re.IGNORECASE):
        none_cb = '[X]'

    # Preserve any content written after ISPM/Codex/OIE checkbox lines
    def _after_option(pattern):
        m = re.search(pattern + r'[^:\n]*:\s*([^\n\[]+)', cell_text, re.IGNORECASE)
        return m.group(1).strip() if m else ''

    codex_extra = _after_option(r'Codex Alimentarius')
    oie_extra   = _after_option(r'World Organ\w+ for Animal Health')
    ippc_extra  = _after_option(r'International Plant Protection')

    # Detect "Is there a relevant international standard? Yes [ ] No [ ]"
    m_yn = re.search(
        r'(?:Is there a relevant international standard|Existe alguna norma).*?(?:Yes|S[íi])\s*(\[[\sXx☒]*\])\s*No\s*(\[[\sXx☒]*\])',
        cell_text, re.IGNORECASE | re.DOTALL,
    )
    has_std_yes = '[X]' if m_yn and m_yn.group(1).strip('[] ').lower() in ('x', '☒') else '[  ]'
    has_std_no  = '[X]' if m_yn and m_yn.group(2).strip('[] ').lower() in ('x', '☒') else '[  ]'

    # Detect "Does this measure conform to the international standard? Yes [ ] No [ ]"
    m_conf = re.search(
        r'(?:Does this measure conform|se ajusta).*?(?:Yes|S[íi])\s*(\[[\sXx☒]*\])\s*No\s*(\[[\sXx☒]*\])',
        cell_text, re.IGNORECASE | re.DOTALL,
    )
    conf_yes = '[X]' if m_conf and m_conf.group(1).strip('[] ').lower() in ('x', '☒') else '[  ]'
    conf_no  = '[X]' if m_conf and m_conf.group(2).strip('[] ').lower() in ('x', '☒') else '[  ]'

    lines = [f'관련 국제기준이 있는가? {has_std_yes} 예  {has_std_no} 아니오  있다면, 해당 기준을 표시']
    lines.append(f'{codex_cb} 국제식품규격위원회(Codex Alimentarius Commission) [예 ; Codex 규정 또는 관련문서의 제목 또는 문서번호] : {codex_extra}')
    lines.append(f'{oie_cb}  세계동물보건기구(OIE) (예 : 육상동물 또는 수생동물 위생규약, Chapter 번호) :  {oie_extra}')
    ippc_label = f'{ippc_cb} 국제식물보호협약(International Plant Protection Convention) [예: 식물위생조치를 위한 국제 기준(ISPM) 번호] :'
    ispm_str = _expand_ispm_numbers(ippc_extra)
    lines.append(f'{ippc_label} {ispm_str or ippc_extra}')
    lines.append(f'{none_cb}  없음')
    lines.append('제안된 규정이 관련 국제기준과 일치하는가?')
    lines.append(f'{conf_yes} 예   {conf_no} 아니오')
    lines.append('그렇지 않다면 가능한 경우는 항상 국제기준과 어떻게 다르고 왜 그러한지 설명:')
    return lines


def _row_other_docs(cell_text, t):
    lines = ['활용 가능한 다른 관련문서 및 언어:']
    doc_kr = t.get('기타문서', '')
    if doc_kr:
        lines.append(doc_kr)
    for url in re.findall(r'https?://\S+', cell_text):
        lines.append(url)
    lang_m = re.search(r'\(available in ([^)]+)\)', cell_text, re.IGNORECASE)
    if lang_m:
        lang_kr = LANG_KR.get(lang_m.group(1).strip().lower(), lang_m.group(1).strip())
        lines.append(f'({lang_kr}로 이용가능)')
    return lines


def _row_adoption_date(cell_text, t):
    def _extract(label_pattern):
        m = re.search(label_pattern + r'[^:]*:\s*(.+?)(?=\n|Proposed date of pub|$)',
                      cell_text, re.IGNORECASE | re.DOTALL)
        return _translate_date_phrase(m.group(1).strip()) if m else '추후 결정'

    adopt = _extract(r'Proposed date of adoption')
    pub   = _extract(r'Proposed date of publication')
    return [
        f'채택예정일 [날짜(일/월/년)]: {adopt}',
        f'공표예정일 [날짜(일/월/년)]: {pub}',
    ]


def _row_entry_force(cell_text, t):
    six_cb   = _checkbox(cell_text, 'Six months')
    trade_cb = _checkbox(cell_text, 'Trade facilitating')

    m = re.search(r'and/or\s*\(dd/mm/yy\):\s*(.+?)(?:\n|$)', cell_text, re.IGNORECASE)
    date_kr = _translate_date_phrase(m.group(1).strip()) if m else ''

    return [
        f'효력 발생일: {six_cb} 공표일로부터6개월 후, 그리고/또는 [날짜(일/월/년)]: {date_kr}',
        f'{trade_cb}  무역 원활화 조치',
    ]


def _row_comments(cell_text, t):
    sixty_cb = _checkbox(cell_text, 'Sixty days')
    nna_cb   = _checkbox(cell_text, 'National Notification Authority')
    neq_cb   = _checkbox(cell_text, 'National Enquiry Point')

    m = re.search(r'and/or\s*\(dd/mm/yy\):\s*(.+?)(?:\n|Agency|$)',
                  cell_text, re.IGNORECASE | re.DOTALL)
    date_kr = _translate_date_phrase(m.group(1).strip()) if m else ''

    # Email/address/phone already exist in the original cell — not duplicated here
    return [
        f'의견제출 마감일: {sixty_cb} 통보문 배포일로부터 60일 후 그리고/또는 [날짜(일/월/년)]: {date_kr}',
        f'의견 처리 담당기관 또는 관계당국: {nna_cb} 국가 통보처, {neq_cb} 국가 문의처 또는 (존재할 경우) 타 기관의 주소, 팩스 번호, 이메일 주소:',
    ]


def _row_texts_available(cell_text, t):
    nna_cb = _checkbox(cell_text, 'National Notification Authority')
    neq_cb = _checkbox(cell_text, 'National Enquiry Point')
    # Contact details already exist in the original cell — not duplicated here
    return [f'전문 입수가 가능한 곳: {nna_cb} 국가 통보처, {neq_cb} 국가 문의처 또는 (존재할 경우) 타 기관의 주소, 팩스 번호, 이메일 주소:']


def _row_addendum_intro(cell_text, t):
    date_m = re.search(r'received on\s+(\d{1,2}\s+\w+\s+\d{4})', cell_text, re.IGNORECASE)
    date_kr = _translate_date(date_m.group(1)) if date_m else ''
    country_kr = t.get('통보국_kr', '')
    if date_kr and country_kr:
        return [f'{date_kr}에 수신한 다음 전문은 {country_kr} 대표단의 요청에 따라 회람됨.']
    if country_kr:
        return [f'다음 전문은 {country_kr} 대표단의 요청에 따라 회람됨.']
    return []


def _row_addendum_country_advises(cell_text, t):
    content_kr = t.get('내용', '')
    if not content_kr:
        return []
    return [content_kr]


def _row_addendum_concerns(cell_text, t):
    lines = ['이 추가사항은 다음에 관한 것임：']
    for eng_prefix, kr_label in ADDENDUM_CONCERN_OPTIONS:
        cb = _checkbox(cell_text, eng_prefix)
        lines.append(f'{cb}\t{kr_label}')
    return lines


def _row_addendum_comment_period_sec(cell_text, t):
    sixty_cb = _checkbox(cell_text, 'Sixty days')
    if re.search(r'not applicable', cell_text, re.IGNORECASE):
        date_kr = '해당사항 없음'
    else:
        m = re.search(r'and/or\s*\(dd/mm/yy\):\s*(.+?)(?:\n|$)', cell_text, re.IGNORECASE)
        date_kr = _translate_date_phrase(m.group(1).strip()) if m else ''
    return [
        '의견수렴기간: (추가사항이 제품 및/또는 잠재적으로 영향을 받는 회원국에 관하여 이전에 통보된 조치의 범위를 확대하는 경우, 일반적으로 최소 60일의 추가 의견수렴기간을 제시해야 함. 원래 발표했던 의견수렴 마감일의 연장 등 다른 상황에서는 추가통보문에서 제시되는 의견수렴기간이 달라질 수 있음. )',
        f'{sixty_cb}        추가통보문 배포일로부터 60일 및/또는 (일/월/년): {date_kr}',
    ]


def _row_addendum_agency_comments(cell_text, t):
    nna_cb = _checkbox(cell_text, 'National Notification Authority')
    neq_cb = _checkbox(cell_text, 'National Enquiry Point')
    lines = [f'의견처리 담당기관 또는 당국: {nna_cb} 국가통보처, {neq_cb} 국가질의처. 다른 기관의 주소, 팩스번호 및 이메일 주소(있는 경우)：']
    email = _extract_email(cell_text)
    if email:
        lines.append(email)
    return lines


ROW_BUILDERS = {
    'notifying_member': _row_notifying_member,
    'agency':           _row_agency,
    'products':         _row_products,
    'regions':          _row_regions,
    'title':            _row_title,
    'description':      _row_description,
    'objective':        _row_objective,
    'standards':        _row_standards,
    'other_docs':       _row_other_docs,
    'adoption_date':    _row_adoption_date,
    'entry_force':      _row_entry_force,
    'comments':         _row_comments,
    'texts_available':  _row_texts_available,
    # Addendum-specific
    'addendum_intro':              _row_addendum_intro,
    'addendum_country_advises':    _row_addendum_country_advises,
    'addendum_concerns':           _row_addendum_concerns,
    'addendum_comment_period_sec': _row_addendum_comment_period_sec,
    'addendum_agency_comments':    _row_addendum_agency_comments,
}


# ── Main ──────────────────────────────────────────────────────────────────────

_TITLE_KR = {
    'NOTIFICATION OF EMERGENCY MEASURES': '긴급조치 통보문',
    'NOTIFICATION': '통보문',
    'ADDENDUM': '추가',
}

ADDENDUM_SKIP_ROWS = frozenset({'adoption_date'})

ADDENDUM_CONCERN_OPTIONS = [
    ('Modification of final date for comments', '의견수렴 마감일 변경'),
    ('Notification of adoption',                '규정의 채택, 공표 또는 발효 통보'),
    ('Modification of content',                 '이전에 통보한 규정안의 내용 및/또는 범위 변경'),
    ('Withdrawal of proposed regulation',       '규정안 철회'),
    ('Change in proposed date',                 '채택, 공표 또는 발효예정일 변경'),
    ('Other',                                   '기타:'),
]


def _insert_paragraph_after_para(para, text, font_size=None):
    """Insert a new paragraph with Korean text immediately after para using XML.
    Copies paragraph and run properties (indentation, bold, underline, etc.) from the source."""
    import copy
    # Only apply bold/italic/underline when ALL content runs share that property;
    # mixed paragraphs (e.g. bold label + plain body) should produce plain Korean text.
    non_empty_runs = [r for r in para.runs if r.text.strip()]
    if non_empty_runs:
        src_bold      = True if all(r.bold      is True for r in non_empty_runs) else None
        src_italic    = True if all(r.italic    is True for r in non_empty_runs) else None
        src_underline = True if all(r.underline is True for r in non_empty_runs) else None
        if font_size is None:
            font_size = next((r.font.size for r in non_empty_runs if r.font.size), None)
    else:
        src_bold = src_italic = src_underline = None

    new_p = OxmlElement('w:p')
    # Copy paragraph-level formatting from source (indentation, tabs, alignment)
    src_pPr = para._p.find(qn('w:pPr'))
    if src_pPr is not None:
        new_pPr = copy.deepcopy(src_pPr)
        for rpr in new_pPr.findall(qn('w:rPr')):
            new_pPr.remove(rpr)
        new_p.append(new_pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    KOREAN_FONT)
    rFonts.set(qn('w:hAnsi'),   KOREAN_FONT)
    rFonts.set(qn('w:eastAsia'), KOREAN_FONT)
    rPr.append(rFonts)
    if src_bold:
        rPr.append(OxmlElement('w:b'))
    if src_italic:
        rPr.append(OxmlElement('w:i'))
    if src_underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    if font_size:
        sz_val = str(int(font_size / 6350))
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), sz_val)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), sz_val)
        rPr.append(sz)
        rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    if text.startswith(' ') or text.endswith(' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    new_p.append(r)
    para._p.addnext(new_p)


def _translate_doc_titles(doc):
    """
    Find the NOTIFICATION / ADDENDUM paragraphs and append Korean translations.
    Checks both top-level paragraphs (standard docs) and table cells (addendum docs).
    """
    for para in doc.paragraphs:
        text = para.text.strip()
        kr = _TITLE_KR.get(text.upper(), '')
        if not kr:
            continue
        font_size = None
        src_run = None
        for run in para.runs:
            if run.font.size and src_run is None:
                font_size = run.font.size
            if run.text.strip() and src_run is None:
                src_run = run
        run_br = para.add_run()
        br_el = OxmlElement('w:br')
        run_br._r.append(br_el)
        run_kr = para.add_run(kr)
        _apply_korean_font(run_kr)
        if font_size:
            run_kr.font.size = font_size
        if src_run:
            if src_run.bold is not None:
                run_kr.bold = src_run.bold
            if src_run.italic is not None:
                run_kr.italic = src_run.italic
            if src_run.underline is not None:
                run_kr.underline = src_run.underline

    # Also check table cells (addendum docs have no top-level paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in _unique_cells(row):
                cell_text = cell.text.strip()
                kr = _TITLE_KR.get(cell_text.upper(), '')
                if not kr:
                    continue
                font_size, _, bold, italic, underline = _get_cell_style(cell)
                _add_paragraph(cell, kr, font_size, bold=bold, italic=italic, underline=underline)


def _detect_row_type(text: str):
    prefix = text[:150].lower()
    for row_type, patterns in ROW_PATTERNS.items():
        if any(p in prefix for p in patterns):
            return row_type
    return None


def _translate_addendum_reg_title(doc, translations):
    """
    For addendum docs: inject the Korean regulation title into the cell that
    follows the '___' separator line (positional, since the cell has no label).
    In most WTO addendum files the separator is a top-level paragraph and the
    title is the first row of the only table.
    """
    title_kr = translations.get('제목', '')
    if not title_kr:
        return

    # Case 1: ___ is a top-level paragraph (common in WTO addendum docs).
    # The regulation title is then the first non-empty, non-detectable table cell.
    for para in doc.paragraphs:
        if re.match(r'^_+$', para.text.strip()):
            for table in doc.tables:
                for row in table.rows:
                    for cell in _unique_cells(row):
                        ctext = cell.text.strip()
                        if ctext and not _detect_row_type(ctext):
                            _add_paragraph(cell, title_kr, _get_cell_style(cell)[0])
                            return
            return

    # Case 2: ___ is inside a table cell.
    for table in doc.tables:
        cells_flat = [c for row in table.rows for c in _unique_cells(row)]
        for i, cell in enumerate(cells_flat):
            if re.match(r'^_+$', cell.text.strip()):
                for j in range(i + 1, min(i + 4, len(cells_flat))):
                    candidate = cells_flat[j]
                    ctext = candidate.text.strip()
                    if ctext and not _detect_row_type(ctext):
                        _add_paragraph(candidate, title_kr, _get_cell_style(candidate)[0])
                        return


def _translate_addendum_top_paragraphs(doc, translations):
    """Translate top-level paragraphs in addendum docs (intro text lives in doc.paragraphs)."""
    for para in doc.paragraphs:
        row_type = _detect_row_type(para.text)
        if not row_type or row_type not in ROW_BUILDERS:
            continue
        if row_type in ADDENDUM_SKIP_ROWS:
            continue
        korean_lines = ROW_BUILDERS[row_type](para.text, translations)
        if not korean_lines:
            continue
        font_size = next((r.font.size for r in para.runs if r.font.size), None)
        for line in reversed(korean_lines):
            _insert_paragraph_after_para(para, line, font_size)


def create_bilingual_docx(
    source_path: str,
    translations: dict,
    is_non_english: bool = False,
    is_addendum: bool = False,
) -> str:
    """
    Create a bilingual Word file by appending Korean translations to each
    content cell. Returns the output file path (*_번역.docx).
    """
    source = Path(source_path)
    output_path = source.parent / (source.stem + '_번역.docx')
    shutil.copy2(source_path, output_path)

    doc = Document(str(output_path))

    _translate_doc_titles(doc)

    # Translate top-level paragraphs (addendum intro lives in doc.paragraphs, not tables)
    if is_addendum:
        _translate_addendum_top_paragraphs(doc, translations)

    for table in doc.tables:
        rows = list(table.rows)
        for row_idx, row in enumerate(rows):
            cells = _unique_cells(row)
            if not cells:
                continue

            content_cell = cells[-1]
            row_type = _detect_row_type(content_cell.text)

            # Fallback: check non-last cells (layout B: label | content)
            if not row_type and len(cells) > 1:
                for c in cells[:-1]:
                    row_type = _detect_row_type(c.text)
                    if row_type:
                        break

            if not row_type or row_type not in ROW_BUILDERS:
                continue

            if is_addendum and row_type in ADDENDUM_SKIP_ROWS:
                continue

            # Skip standalone checkbox rows in addendum docs — they belong to a section
            # already consumed by the look-ahead in addendum_concerns /
            # addendum_comment_period_sec processing above.
            if is_addendum and re.match(r'^\[[ X\xa0]\]', content_cell.text.strip(), re.IGNORECASE):
                continue

            # For addendum sections spanning multiple rows (each checkbox in its own row),
            # collect ALL row text so checkbox states are detected correctly, and track
            # the last row so the Korean block lands below all English checkboxes.
            section_text = content_cell.text
            target_cell = content_cell
            if is_addendum and row_type in ('addendum_concerns', 'addendum_comment_period_sec'):
                for nri in range(row_idx + 1, len(rows)):
                    ncells = _unique_cells(rows[nri])
                    if ncells:
                        ntext = ncells[-1].text.strip()
                        if re.match(r'^\[[ X\xa0]\]', ntext, re.IGNORECASE):
                            section_text += '\n' + ncells[-1].text
                            target_cell = ncells[-1]
                        else:
                            break

            korean_lines = ROW_BUILDERS[row_type](section_text, translations)
            if not korean_lines:
                continue

            font_size, para_style, bold, italic, underline = _get_cell_style(content_cell)

            if is_addendum and row_type in ('addendum_concerns', 'addendum_comment_period_sec'):
                # Append entire Korean block after the last row of this section
                t_font, _, t_bold, t_italic, t_underline = _get_cell_style(target_cell)
                for line in korean_lines:
                    _add_paragraph(target_cell, line, t_font, para_style,
                                   bold=t_bold, italic=t_italic, underline=t_underline)
            elif is_addendum and row_type == 'addendum_country_advises':
                # Insert Korean right after the matching paragraph (right below body text)
                patterns = ROW_PATTERNS.get(row_type, [])
                matching_para = next(
                    (p for p in content_cell.paragraphs
                     if p.text.strip() and any(pt in p.text[:150].lower() for pt in patterns)),
                    None,
                )
                if matching_para:
                    for line in reversed(korean_lines):
                        _insert_paragraph_after_para(matching_para, line, font_size)
                else:
                    for line in korean_lines:
                        _add_paragraph(content_cell, line, font_size, para_style,
                                       bold=bold, italic=italic, underline=underline)
            elif row_type == 'comments':
                # 의견제출 마감일 → right after the date/sixty-days paragraph
                # 의견 처리 담당기관 → right after the agency/"other body:" paragraph
                # Contact details already exist in the original cell — not touched
                date_para = next(
                    (p for p in content_cell.paragraphs if re.search(
                        r'final date for comments|sixty days|and/or \(dd/mm/yy\)',
                        p.text, re.IGNORECASE)),
                    None,
                )
                agency_para = next(
                    (p for p in content_cell.paragraphs if re.search(
                        r'agency or authority|national notification authority|national enquiry point|other body',
                        p.text, re.IGNORECASE)),
                    None,
                )
                if korean_lines:
                    if date_para:
                        _insert_paragraph_after_para(date_para, korean_lines[0], font_size)
                    else:
                        _add_paragraph(content_cell, korean_lines[0], font_size, para_style,
                                       bold=bold, italic=italic, underline=underline)
                if len(korean_lines) > 1:
                    if agency_para:
                        _insert_paragraph_after_para(agency_para, korean_lines[1], font_size)
                    else:
                        _add_paragraph(content_cell, korean_lines[1], font_size, para_style,
                                       bold=bold, italic=italic, underline=underline)
            elif row_type == 'texts_available':
                # 전문 입수가 가능한 곳 → right after the "Text(s) available from:" paragraph
                # Contact details already exist in the original cell — not touched
                anchor_para = next(
                    (p for p in content_cell.paragraphs if re.search(
                        r'texts? available from', p.text, re.IGNORECASE)),
                    None,
                )
                if korean_lines:
                    if anchor_para:
                        _insert_paragraph_after_para(anchor_para, korean_lines[0], font_size)
                    else:
                        _add_paragraph(content_cell, korean_lines[0], font_size, para_style,
                                       bold=bold, italic=italic, underline=underline)
            else:
                for line in korean_lines:
                    _add_paragraph(content_cell, line, font_size, para_style,
                                   bold=bold, italic=italic, underline=underline)

            if is_non_english and row_type in ('title', 'description'):
                _set_cell_bg(content_cell, LIME_RGB)

    if is_addendum:
        _translate_addendum_reg_title(doc, translations)

    doc.save(str(output_path))
    return str(output_path)
