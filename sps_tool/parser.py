"""
WTO SPS notification Word document parser.
Extracts structured fields from the standardized WTO SPS form.
"""
import re
import os
from pathlib import Path
import docx
from docx.oxml.ns import qn


# ── Label patterns for each field (English and Korean variants) ──────────────
LABEL_PATTERNS = {
    'notifying_member': ['notifying member', '통보회원국', 'notifying country',
                         'miembro que notifica'],
    'agency':           ['agency responsible', '담당기관', 'responsible agency',
                         'organismo responsable'],
    'products':         ['products covered', '적용대상품목', 'products',
                         'productos'],
    'regions':          ['regions/countries', 'regions', 'countries', '지역/국가', '국가/지역',
                         'regiones'],
    'title':            ['title', '제목',
                         'título', 'titulo'],
    'description':      ['description of content', '내용', 'description',
                         'descripción del contenido', 'descripcion del contenido'],
    'objective':        ['objective', '목적', 'reason',
                         'objetivo'],
    'standards':        ['international standards', '국제기준', 'international standard',
                         'existe una norma'],
    'adoption_date':    ['proposed date of adoption', 'date of adoption', '채택예정일',
                         'proposed date of publication', '발행예정일',
                         'fecha propuesta de adopci'],
    'comment_deadline': ['final date for comments', '의견마감일', 'comment period', 'date for comments',
                         'fecha límite', 'fecha limite'],
    'entry_force':      ['proposed date of entry into force', '발효예정일', 'entry into force',
                         'entrada en vigor'],
    'distribution':     ['distribution date', '배포일', 'circulated', 'circulation date'],
    'other_docs':       ['other relevant documents', '활용 가능한 다른 관련문서',
                         'otros documentos'],
}

OBJECTIVE_MAP = {
    # English
    'food safety':        '식품안전',
    'animal health':      '동물위생',
    'plant protection':   '식물보호',
    'protect humans':     '동식물 병해충 또는 질병으로부터 사람 보호',
    'protect territory':  '해충으로 인한 피해로부터 영토 보호',
    'protect humans from animal': '동식물 병해충 또는 질병으로부터 사람 보호',
    # Spanish
    'inocuidad de los alimentos':       '식품안전',
    'sanidad animal':                   '동물위생',
    'preservación de los vegetales':    '식물보호',
    'preservacion de los vegetales':    '식물보호',
    'protección de la salud humana':    '동식물 병해충 또는 질병으로부터 사람 보호',
    'proteccion de la salud humana':    '동식물 병해충 또는 질병으로부터 사람 보호',
    'protección del territorio':        '해충으로 인한 피해로부터 영토 보호',
    'proteccion del territorio':        '해충으로 인한 피해로부터 영토 보호',
}

DOC_NUMBER_RE = re.compile(
    r'G/SPS/[A-Z]+/[A-Z]{2,3}/[\d]+(?:/Add\.[\d]+)?',
    re.IGNORECASE
)


def _unique_cells(row):
    """Return deduplicated cells from a table row (merged cells repeat in python-docx)."""
    seen = set()
    result = []
    for cell in row.cells:
        cid = id(cell._tc)
        if cid not in seen:
            seen.add(cid)
            result.append(cell)
    return result


def _cell_text(cell):
    return cell.text.strip()


def _all_text(doc):
    """Get all text from paragraphs + tables for pattern searching."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in _unique_cells(row):
                parts.append(cell.text)
    return '\n'.join(parts)


def _extract_doc_number(text, filename=''):
    """Find the WTO document symbol in text, fall back to filename parsing."""
    # Filename takes priority for standard WTO SPS filenames — the document body
    # often references the base number before the addendum number, causing the
    # regex to capture the wrong (shorter) form.
    base = Path(filename).stem.upper()
    base = re.sub(r'_번역$', '', base)
    m2 = re.match(r'GSPSE?N([A-Z]{2,3})(\d+)(A(\d+))?$', base)
    if m2:
        country = m2.group(1)
        number  = m2.group(2)
        add_num = m2.group(4)
        result  = f'G/SPS/N/{country}/{number}'
        if add_num:
            result += f'/Add.{add_num}'
        return result

    m = DOC_NUMBER_RE.search(text)
    if m:
        return m.group().upper()

    return ''


def _detect_type(full_text, doc_number, filename=''):
    """Return dict with is_emergency and is_addendum flags."""
    head = full_text[:600].lower()
    is_emergency = (
        'emergency' in head or
        'g/sps/n/ems' in doc_number.lower()
    )
    is_addendum = (
        'addendum' in head or
        '/add.' in doc_number.lower() or
        re.search(r'A\d+$', Path(filename).stem.upper()) is not None
    )
    return {'is_emergency': is_emergency, 'is_addendum': is_addendum}


def _match_label(cell_text, patterns):
    t = cell_text.lower()
    return any(p in t for p in patterns)


def _extract_field_from_tables(doc, label_patterns):
    """
    Find the content for a field by its label patterns.

    Layout A (WTO standard): ['1.', 'Label: content...']
      The label is embedded at the start of the last cell.
      Extract everything after the first colon.

    Layout B (older format): ['Label cell', 'Content cell']
      The label is in a dedicated earlier cell; return the last cell.
    """
    for table in doc.tables:
        for row in table.rows:
            cells = _unique_cells(row)
            if len(cells) < 2:
                continue

            content_cell = cells[-1]
            content_text = _cell_text(content_cell)
            if not content_text:
                continue

            # Layout A: label at the start of the content cell (first line)
            first_line = content_text.split('\n')[0][:150]
            if _match_label(first_line, label_patterns):
                colon_pos = content_text.find(':')
                if colon_pos != -1:
                    return content_text[colon_pos + 1:].strip()
                return content_text

            # Layout B: label is in a dedicated earlier cell
            for cell in cells[:-1]:
                if _match_label(_cell_text(cell), label_patterns):
                    if content_text and len(content_text) > 1 and content_text != _cell_text(cell):
                        return content_text
    return ''


# Fields extracted in one pass — covers all fields called individually in parse_notification
_BULK_EXTRACT_FIELDS = [
    'notifying_member', 'agency', 'products', 'regions',
    'title', 'description', 'objective', 'other_docs',
    'comment_deadline', 'entry_force', 'adoption_date',
]


def _extract_description_paragraphs(doc) -> str:
    """
    Extract the description field paragraph by paragraph, stripping the label
    from the first paragraph (Layout A).  Returns paragraphs joined by \\n.
    Falls back to '' so the caller can use plain-text extraction.
    """
    desc_patterns = LABEL_PATTERNS['description']
    for table in doc.tables:
        for row in table.rows:
            cells = _unique_cells(row)
            if len(cells) < 2:
                continue
            content_cell = cells[-1]
            content_text = _cell_text(content_cell)
            if not content_text:
                continue

            first_line = content_text.split('\n')[0][:150].lower()
            is_layout_a = any(p in first_line for p in desc_patterns)
            is_layout_b = False
            if not is_layout_a:
                for cell in cells[:-1]:
                    if _match_label(_cell_text(cell), desc_patterns):
                        is_layout_b = True
                        break

            if not (is_layout_a or is_layout_b):
                continue

            lines = []
            strip_label = is_layout_a  # first paragraph of Layout A has "Description of content:" prefix
            for para in content_cell.paragraphs:
                text = para.text
                if strip_label:
                    ci = text.find(':')
                    text = text[ci + 1:].strip() if ci != -1 else text
                    strip_label = False
                else:
                    text = text.strip()
                if text:
                    lines.append(text)

            return '\n'.join(lines)

    return ''


def _extract_all_fields(doc) -> dict:
    """
    Single table-scan extraction for all labeled fields.
    Returns a dict keyed by field name (same keys as LABEL_PATTERNS).
    Replaces 11 individual _extract_field_from_tables calls with one pass.
    """
    label_map = {k: LABEL_PATTERNS[k] for k in _BULK_EXTRACT_FIELDS}
    results = {k: '' for k in label_map}
    remaining = set(label_map.keys())

    for table in doc.tables:
        if not remaining:
            break
        for row in table.rows:
            if not remaining:
                break
            cells = _unique_cells(row)
            if len(cells) < 2:
                continue
            content_cell = cells[-1]
            content_text = _cell_text(content_cell)
            if not content_text:
                continue

            first_lower = content_text.split('\n')[0][:150].lower()
            found = set()

            for field_name in list(remaining):
                patterns = label_map[field_name]
                # Layout A: label embedded at start of last cell
                if any(p in first_lower for p in patterns):
                    colon_pos = content_text.find(':')
                    results[field_name] = (
                        content_text[colon_pos + 1:].strip() if colon_pos != -1
                        else content_text
                    )
                    found.add(field_name)
                    continue
                # Layout B: label in an earlier cell
                for cell in cells[:-1]:
                    ct = _cell_text(cell)
                    if _match_label(ct, patterns):
                        if content_text and len(content_text) > 1 and content_text != ct:
                            results[field_name] = content_text
                            found.add(field_name)
                        break

            remaining -= found

    return results


def _extract_objectives(full_text: str) -> list:
    """
    Find checked objectives ([X] or ☒ markers) and return Korean phrases.
    Searches the full document text for a checked mark immediately before each
    objective key, so it works correctly whether options are in separate cells
    (English layout) or all in one cell (Spanish/Portuguese layout).
    """
    checked = []
    for key, kor_val in OBJECTIVE_MAP.items():
        pattern = r'(?:\[x\]|☒)\s*' + re.escape(key)
        if re.search(pattern, full_text, re.IGNORECASE) and kor_val not in checked:
            checked.append(kor_val)
    return checked


def _extract_regions(regions_raw: str, full_text: str) -> str:
    """
    Determine final regions value from pre-extracted raw field and full text.
    Returns '모든 교역국' if all trading partners is checked, otherwise
    returns specific country names or the raw fallback.
    Handles English and Spanish WTO notification forms.
    """
    # English / Korean: all trading partners
    if re.search(r'\[x\].*?all trading partners', full_text, re.IGNORECASE | re.DOTALL):
        return '모든 교역국'
    if re.search(r'\[x\].*?모든 교역국', full_text, re.DOTALL):
        return '모든 교역국'
    # Spanish: "Todos los interlocutores comerciales"
    if re.search(r'\[x\].*?todos los interlocutores', full_text, re.IGNORECASE | re.DOTALL):
        return '모든 교역국'

    # English: specific regions
    specific_match = re.search(
        r'\[x\][^\n]*?specific regions?(?:\s+or\s+countries?)?\s*:\s*([^\n\[]+)',
        full_text, re.IGNORECASE
    )
    if specific_match:
        return specific_match.group(1).strip()

    # Spanish: "Regiones o países específicos: [country]"
    specific_match = re.search(
        r'\[x\][^\n]*?espec[íi]ficos\s*:\s*([^\n\[]+)',
        full_text, re.IGNORECASE
    )
    if specific_match:
        return specific_match.group(1).strip()

    return regions_raw


def _detect_language(text):
    """
    Detect dominant source language from character distribution.
    Returns 'en', 'es', or 'pt'.
    """
    if not text:
        return 'en'
    # Spanish/Portuguese indicator characters
    sp_pt_chars = set('áéíóúàèìòùâêîôûãõñüçÁÉÍÓÚÀÈÌÒÙÂÊÎÔÛÃÕÑÜÇ')
    count = sum(1 for c in text if c in sp_pt_chars)
    if count == 0:
        return 'en'
    # Very rough heuristic: ã/õ = probably Portuguese, ñ = probably Spanish
    pt_chars = set('ãõÃÕ')
    es_chars = set('ñÑ')
    pt_count = sum(1 for c in text if c in pt_chars)
    es_count = sum(1 for c in text if c in es_chars)
    if pt_count > es_count:
        return 'pt'
    return 'es'


def _extract_addendum_fields(doc, full_text):
    """
    For addendum documents, extract the specific addendum metadata.
    Returns a dict with addendum_type and relevant content.
    """
    result = {
        'addendum_concerns': [],
        'addendum_country': '',
        'addendum_received_date': '',
        'addendum_content': '',
        'addendum_reg_title': '',
        'addendum_country_advises': '',
    }

    # Extract circulation country from opening sentence
    m = re.search(
        r'being circulated at the request of the delegation of ([A-Z][A-Z ]+)',
        full_text, re.IGNORECASE
    )
    if m:
        result['addendum_country'] = m.group(1).strip()

    # Extract received date
    m2 = re.search(
        r'received on ([0-9\w ,]+)',
        full_text, re.IGNORECASE
    )
    if m2:
        result['addendum_received_date'] = m2.group(1).strip()

    # Extract regulation title and country advises paragraph from body
    sep_pos = full_text.find('___')
    if sep_pos != -1:
        after_sep = full_text[sep_pos:].lstrip('_').lstrip()
        lines = [l.strip() for l in after_sep.split('\n') if l.strip()]
        if lines:
            result['addendum_reg_title'] = lines[0]
        advises_m = re.search(
            r'\w[\w\s]+ hereby advises.+?(?=\n\s*\n|\Z)',
            after_sep, re.DOTALL | re.IGNORECASE,
        )
        if advises_m:
            result['addendum_country_advises'] = advises_m.group().strip()

    # Find checked addendum type boxes
    addendum_types = {
        'notification of adoption':  '채택·발행·발효 통보',
        'modification of final date': '의견마감일 변경',
        'modification of content':   '내용/범위 변경',
        'withdrawal':                '규정 철회',
        'change in proposed dates':  '제안 일자 변경',
    }
    for eng, kor in addendum_types.items():
        pattern = r'\[x\][^\n]*' + re.escape(eng)
        if re.search(pattern, full_text, re.IGNORECASE):
            result['addendum_concerns'].append(kor)

    return result


def parse_notification(docx_path: str) -> dict:
    """
    Parse a WTO SPS notification Word file and return structured fields.

    Returns a dict with all extracted raw fields. The LLM will later
    translate and normalize these into Korean institutional language.
    """
    doc = docx.Document(docx_path)
    filename = os.path.basename(docx_path)
    full_text = _all_text(doc)

    result = {
        'filename':             filename,
        'doc_number':           '',
        'is_emergency':         False,
        'is_addendum':          False,
        'notifying_member':     '',
        'agency':               '',
        'products':             '',
        'regions':              '',
        'title':                '',
        'description':          '',
        'objectives_raw':       [],
        'objectives_korean':    [],
        'comment_deadline_raw': '',
        'entry_force_raw':      '',
        'adoption_date_raw':    '',
        'source_language':      'en',
        'other_docs':           '',
        'objective_text':       '',
        'addendum':             {},
    }

    # ── Document number ────────────────────────────────────────────────────
    result['doc_number'] = _extract_doc_number(full_text, filename)

    # ── Notification type ──────────────────────────────────────────────────
    type_flags = _detect_type(full_text, result['doc_number'], filename)
    result['is_emergency'] = type_flags['is_emergency']
    result['is_addendum']  = type_flags['is_addendum']

    # ── Field extraction (single pass over all tables) ─────────────────────
    extracted = _extract_all_fields(doc)
    # Keep only the first non-empty line of the notifying member cell.
    # The WTO form cell often continues with "If applicable, name of local government..."
    # which is a boilerplate label (not a value) and confuses downstream LLM processing.
    raw_member = extracted['notifying_member']
    result['notifying_member'] = next(
        (ln.strip() for ln in raw_member.split('\n') if ln.strip()), ''
    )
    result['agency']               = extracted['agency']
    result['products']             = extracted['products']
    result['title']                = extracted['title']
    result['description']          = _extract_description_paragraphs(doc) or extracted['description']
    result['objective_text']       = extracted['objective']
    result['other_docs']           = extracted['other_docs']
    result['comment_deadline_raw'] = extracted['comment_deadline']
    result['entry_force_raw']      = extracted['entry_force']
    result['adoption_date_raw']    = extracted['adoption_date']

    # ── Regions and objectives use the already-computed full_text ──────────
    result['regions'] = _extract_regions(extracted['regions'], full_text)

    # ── Objectives (checkboxes) ────────────────────────────────────────────
    result['objectives_korean'] = _extract_objectives(full_text)

    # ── Language detection ────────────────────────────────────────────────
    detect_text = result['description'] or result['title'] or result['products']
    result['source_language'] = _detect_language(detect_text)

    # ── Addendum-specific fields ──────────────────────────────────────────
    if result['is_addendum']:
        result['addendum'] = _extract_addendum_fields(doc, full_text)
        body = result['addendum']
        if body.get('addendum_reg_title'):
            result['title'] = body['addendum_reg_title']  # always use body title for addendum
        if body.get('addendum_country_advises') and not result['description']:
            result['description'] = body['addendum_country_advises']

    return result
